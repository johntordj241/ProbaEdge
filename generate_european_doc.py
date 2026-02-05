import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

warnings.filterwarnings("ignore")

# Charger les données
df = pd.read_csv("data/prediction_dataset_enriched_v2.csv")

# Convertir les dates
df["fixture_date"] = pd.to_datetime(df["fixture_date"], utc=True, errors="coerce")
df = df.dropna(subset=["fixture_date"])

# Filtrer LDC (3) et Europa (4)
european = df[df["league_id"].isin([3.0, 4.0])].copy()
european = european.drop_duplicates(subset=["fixture_id"]).sort_values("fixture_date")

# Créer le document Word
doc = Document()
doc.add_heading("📅 Matchs Européens - Prédictions Complètes", 0)
doc.add_paragraph(f"Généré le: 19 janvier 2026")
doc.add_paragraph(f"Total matchs: {len(european)} matchs")
doc.add_paragraph()

# Grouper par ligue
for league_id in [3.0, 4.0]:
    league_name = "🏆 LIGUE DES CHAMPIONS" if league_id == 3.0 else "🎯 EUROPA LEAGUE"
    matches = european[european["league_id"] == league_id].copy()

    if len(matches) == 0:
        continue

    # Titre de la ligue
    doc.add_heading(league_name, level=1)
    doc.add_paragraph(f"{len(matches)} matchs")

    # Grouper par date
    matches["date_str"] = matches["fixture_date"].dt.strftime("%d/%m/%Y")
    dates = sorted(matches["date_str"].unique())

    for date in dates:
        matches_on_date = matches[matches["date_str"] == date].sort_values(
            "fixture_date"
        )

        # Sous-titre date
        doc.add_heading(f"📍 {date}", level=2)

        # Créer une table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"

        # En-têtes
        header_cells = table.rows[0].cells
        header_cells[0].text = "Heure"
        header_cells[1].text = "Matchs"
        header_cells[2].text = "1X2 Probs"
        header_cells[3].text = "Over/Under"
        header_cells[4].text = "Main Pick"

        # Remplir les lignes
        for idx, match in matches_on_date.iterrows():
            time = match["fixture_date"].strftime("%H:%M")

            # Nouvelle ligne
            row_cells = table.add_row().cells

            # Heure
            row_cells[0].text = time

            # Matchs
            row_cells[1].text = f"{match['home_team']}\nvs\n{match['away_team']}"

            # 1X2
            h = match["prob_home"] * 100
            d = match["prob_draw"] * 100
            a = match["prob_away"] * 100
            row_cells[2].text = f"1: {h:.0f}%\nX: {d:.0f}%\n2: {a:.0f}%"

            # Over/Under
            o25 = match["prob_over_2_5"] * 100
            u25 = match["prob_under_2_5"] * 100
            row_cells[3].text = f"O2.5: {o25:.0f}%\nU2.5: {u25:.0f}%"

            # Main Pick
            row_cells[4].text = (
                f"{match['main_pick']}\n({match['main_confidence']*100:.0f}%)"
            )

        doc.add_paragraph()

# Ajouter un résumé
doc.add_page_break()
doc.add_heading("📊 Résumé des Recommandations", level=1)

doc.add_paragraph("Meilleurs Over 2.5:")
over_picks = european[
    european["main_pick"].str.contains("Over", case=False, na=False)
].nlargest(5, "prob_over_2_5")
for idx, match in over_picks.iterrows():
    doc.add_paragraph(
        f"• {match['home_team']} vs {match['away_team']} - Over 2.5 à {match['prob_over_2_5']*100:.0f}% ({match['fixture_date'].strftime('%d/%m %H:%M')})",
        style="List Bullet",
    )

doc.add_paragraph()
doc.add_paragraph("Meilleures victoires dominantes (Home >80%):")
home_strong = european[european["prob_home"] > 0.80].nlargest(5, "prob_home")
for idx, match in home_strong.iterrows():
    doc.add_paragraph(
        f"• {match['home_team']} vs {match['away_team']} - Home à {match['prob_home']*100:.0f}% ({match['fixture_date'].strftime('%d/%m %H:%M')})",
        style="List Bullet",
    )

# Sauvegarder
output_file = "European_Matchs_Predictions.docx"
doc.save(output_file)

print(f"✅ Document créé: {output_file}")
print(f"📊 {len(european)} matchs analysés")
